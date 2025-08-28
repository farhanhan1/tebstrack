"""
Document processing utilities for knowledge base integration
"""

import os
import fitz  # PyMuPDF for PDF
import docx  # python-docx for DOCX
import logging
from typing import Optional

class DocumentLoader:
    @staticmethod
    def load_knowledge_document(file_path: str) -> Optional[str]:
        """
        Load and extract text from PDF or DOCX knowledge documents
        """
        if not os.path.exists(file_path):
            logging.warning(f"Knowledge document not found: {file_path}")
            return None
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return DocumentLoader._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                return DocumentLoader._extract_docx_text(file_path)
            else:
                logging.error(f"Unsupported document format: {file_ext}")
                return None
                
        except Exception as e:
            logging.error(f"Failed to load document {file_path}: {e}")
            return None
    
    @staticmethod
    def _extract_pdf_text(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text.strip()
    
    @staticmethod
    def _extract_docx_text(file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
                    
        return '\n'.join(text).strip()

    @staticmethod
    def chunk_text(text: str, max_chunk_size: int = 3000) -> list:
        """
        Split large text into chunks for better AI processing
        """
        if len(text) <= max_chunk_size:
            return [text]
            
        chunks = []
        sentences = text.split('. ')
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk + sentence) <= max_chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
                
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        return chunks
